"""
계약서 자동 생성 엔진

지원 유형:
  - 자재납품 계약서 (16개 조항, 오티엘 템플릿 기반)
  - 공사 계약서 (26개 조항, 호림ENG 템플릿 기반)

사용:
  from src.contracts.contract_generator import generate_material_contract, generate_construction_contract
  path = generate_material_contract(params, "output/계약서.docx")
"""

import pathlib
import shutil
import logging
from copy import deepcopy
from datetime import datetime

from docx import Document

logger = logging.getLogger(__name__)

# ── 경로 ────────────────────────────────────────────────────────────────────
CONTRACTS_DIR = pathlib.Path(__file__).parent
MATERIAL_TEMPLATE = CONTRACTS_DIR / "자재납품_template.docx"
CONSTRUCTION_TEMPLATE = CONTRACTS_DIR / "공사_template.docx"

# ── 갑(글로우서울) 고정 정보 ─────────────────────────────────────────────────
GAP_INFO = {
    "display": "㈜글로우서울",
    "full": "주식회사 글로우서울",
    "rep": "윤  성  혁",
    "addr_body": "서울특별시 용산구 이태원로17길 23-8, 23-13, 23-17",
    "addr_sign": "서울특별시 용산구 이태원동 164-2",
    "biz_no": "451-88-01049",
}

# 공사 현장 주소 (고정)
SITE_ADDRESS = "서울특별시 종로구 청계천로 93"
PROJECT_NAME = "종로 오블리브 의원 프로젝트"


# ── 한글 금액 변환 ──────────────────────────────────────────────────────────

def num_to_korean_won(amount: int) -> str:
    """정수 금액 → 한글 표기 (원정 포함)
    예: 47_500_000 → '사천칠백오십만원정'
    """
    if amount == 0:
        return "영원정"

    ones = ["", "일", "이", "삼", "사", "오", "육", "칠", "팔", "구"]
    pos = ["", "십", "백", "천"]

    def convert_4(n: int) -> str:
        if n == 0:
            return ""
        result = ""
        for i in range(3, -1, -1):
            d = (n // (10 ** i)) % 10
            if d == 0:
                continue
            result += ones[d] + pos[i]
        return result

    parts = []
    jo = amount // 10 ** 12
    amount %= 10 ** 12
    eok = amount // 10 ** 8
    amount %= 10 ** 8
    man = amount // 10_000
    rest = amount % 10_000

    if jo:
        parts.append(convert_4(jo) + "조")
    if eok:
        parts.append(convert_4(eok) + "억")
    if man:
        parts.append(convert_4(man) + "만")
    if rest:
        parts.append(convert_4(rest))

    return "".join(parts) + "원정"


def format_date_kor(date_str: str) -> str:
    """'2026-02-05' → '2026년 02월 05일'"""
    try:
        d = datetime.strptime(date_str.strip(), "%Y-%m-%d")
        return f"{d.year}년 {d.month:02d}월 {d.day:02d}일"
    except Exception:
        return date_str


def format_date_kor_short(date_str: str) -> str:
    """'2026-02-28' → '2026년 2월 28일'"""
    try:
        d = datetime.strptime(date_str.strip(), "%Y-%m-%d")
        return f"{d.year}년 {d.month}월 {d.day}일"
    except Exception:
        return date_str


# ── DOCX 텍스트 치환 ─────────────────────────────────────────────────────────

def _replace_para(para, replacements: dict) -> bool:
    """단락 runs를 합쳐 치환 후 첫 run에 적용. 변경 시 True 반환."""
    full = "".join(r.text for r in para.runs)
    new = full
    for old, val in replacements.items():
        new = new.replace(old, val)
    if new != full and para.runs:
        para.runs[0].text = new
        for r in para.runs[1:]:
            r.text = ""
        return True
    return False


def replace_in_doc(doc: Document, replacements: dict) -> int:
    """doc 전체(본문 + 표)에 replacements 일괄 적용. 변경 횟수 반환."""
    count = 0
    for para in doc.paragraphs:
        if _replace_para(para, replacements):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_para(para, replacements):
                        count += 1
    return count


# ── 자재납품 계약서 ──────────────────────────────────────────────────────────

def generate_material_contract(params: dict, output_path: str) -> str:
    """
    자재납품 계약서 생성.

    params:
        sub_name        협력사 이름 (예: "오티엘", "유니밋")
        sub_name_full   법인 전체명 (예: "주식회사 오티엘"). 없으면 sub_name 사용
        sub_rep         대표자 (예: "이 봉 식")
        sub_addr        주소 (예: "서울특별시 도봉구 도봉산3길 46, 지하층(도봉동)")
        sub_biz_no      사업자등록번호 (예: "390-86-03578")
        material_type   자재 유형 (예: "조명", "커튼", "위생기구")
        material_desc   물품 상세 (예: "조명(등기구) 제작 및 납품 (견적서 참조)")
        total_amount    계약금액 (int, VAT 별도, 예: 47500000)
        deposit_pct     계약금 비율 (int, 예: 40)
        progress_pct    중도금 비율 (int, 예: 50)  ← 잔금은 자동(10%)
        delivery_date   납품일자 (예: "2026-02-28")
        install_date    설치완료일 (예: "2026-03-16"). 없으면 delivery_date 사용
        contract_date   계약일 (예: "2026-02-03"). 없으면 오늘

    Returns: 생성된 파일 경로
    """
    p = params
    sub_name = p.get("sub_name", "")
    sub_name_full = p.get("sub_name_full") or sub_name
    sub_rep = p.get("sub_rep", "")
    sub_addr = p.get("sub_addr", "")
    sub_biz_no = p.get("sub_biz_no", "")
    mat_type = p.get("material_type", "")
    mat_desc = p.get("material_desc", f"{mat_type} 납품 (견적서 참조)")
    total = int(p.get("total_amount", 0))
    dep_pct = int(p.get("deposit_pct", 40))
    prog_pct = int(p.get("progress_pct", 50))
    bal_pct = 100 - dep_pct - prog_pct  # 잔금
    dep_amt = round(total * dep_pct / 100)
    prog_amt = round(total * prog_pct / 100)
    bal_amt = total - dep_amt - prog_amt

    delivery_raw = p.get("delivery_date", "")
    install_raw = p.get("install_date", "") or delivery_raw
    contract_raw = p.get("contract_date", datetime.today().strftime("%Y-%m-%d"))

    delivery_kor = format_date_kor_short(delivery_raw) if delivery_raw else ""
    install_kor = format_date_kor_short(install_raw) if install_raw else ""
    contract_kor = format_date_kor(contract_raw)

    # 계약일 분리 (연/월/일)
    try:
        cd = datetime.strptime(contract_raw.strip(), "%Y-%m-%d")
        contract_sign = f"{cd.year}년 {cd.month:02d}월 {cd.day:02d}일"
    except Exception:
        contract_sign = contract_raw

    # ── 치환 맵 ──
    replacements = {
        # 제목
        "자재납품(조명) 계약서": f"자재납품({mat_type}) 계약서",
        # 1조
        "자재(조명)공급에 관한": f"자재({mat_type})공급에 관한",
        '㈜글로우서울(이하 \u201c갑\u201d이라 한다)과 ㈜오티엘': f'㈜글로우서울(이하 \u201c갑\u201d이라 한다)과 {sub_name}',
        '㈜오티엘 (이하 \u201c을\u201d이라 한다)': f'{sub_name} (이하 \u201c을\u201d이라 한다)',
        # 4조 대금 지급
        f"계약금(40%) : 일금 일천구백만원정 (￦19,000,000) / 계약체결 이후":
            f"계약금({dep_pct}%) : 일금 {num_to_korean_won(dep_amt)} (￦{dep_amt:,}) / 계약체결 이후",
        f"중도금(50%) : 일금 이천삼백칠십오만원정 (￦23,750,000) / 납품 완료 이후":
            f"중도금({prog_pct}%) : 일금 {num_to_korean_won(prog_amt)} (￦{prog_amt:,}) / 납품 완료 이후",
        "잔\t금(10%) : 일금 사백칠십오만원정 (￦4,750,000) / 하도급 정산 완료 이후":
            f"잔\t금({bal_pct}%) : 일금 {num_to_korean_won(bal_amt)} (￦{bal_amt:,}) / 하도급 정산 완료 이후",
        # 표1 물품
        "조명(등기구) 제작 및 납품 (견적서 참조)": mat_desc,
        # 표1 계약금액
        "일금 : 사천칠백오십만원정 (￦47,500,000) 부가세 별도":
            f"일금 : {num_to_korean_won(total)} (￦{total:,}) 부가세 별도",
        # 표1 납품일자
        "2026년 2월 28일 이내 ( 3월 16일 이전 설치 완료 )":
            f"{delivery_kor} 이내 ( {install_kor} 이전 설치 완료 )",
        # 첨부서류
        "1. 조명 제품사양서 1부": f"1. {mat_type} 제품사양서 1부",
        # 계약일
        "2026년 02월 03일": contract_sign,
        # 갑 서명
        "서울특별시 용산구 이태원로17길 23-8, 23-13, 23-17 사업자등록번호 :		451-88-01049":
            f"{GAP_INFO['addr_body']} 사업자등록번호 :		{GAP_INFO['biz_no']}",
        # 을 서명
        '"을" 주식회사 오티엘': f'"을" {sub_name_full}',
        "주식회사 오티엘": sub_name_full,
        "서울특별시 도봉구 도봉산3길 46, 지하층(도봉동)": sub_addr,
        "390-86-03578": sub_biz_no,
        "이 봉 식": sub_rep,
    }

    doc = Document(MATERIAL_TEMPLATE)
    changed = replace_in_doc(doc, replacements)
    logger.info(f"자재납품 계약서: {changed}개 항목 치환 → {output_path}")

    out = pathlib.Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ── 공사 계약서 ──────────────────────────────────────────────────────────────

def generate_construction_contract(params: dict, output_path: str) -> str:
    """
    공사 계약서 생성.

    params:
        sub_name            협력사명 (예: "호림ENG", "아트음향")
        sub_rep             대표자 (예: "이 만 식")
        sub_addr            주소 (예: "서울특별시 강동구 구천면로 361...")
        construction_type   공사 유형 (예: "위생배관공사", "음향공사", "전기공사")
        supply_amount       공급가액 (VAT 제외, int, 예: 48000000)
        start_date          착공일 (예: "2026-02-05")
        end_date            준공예정일 (예: "2026-03-16")
        progress2_cond      중도금2차 조건 (예: "공종 작업 완료 이후", "시운전 완료 이후")
        contract_date       계약일 (예: "2026-02-05"). 없으면 오늘

    Returns: 생성된 파일 경로
    """
    p = params
    sub_name = p.get("sub_name", "")
    sub_rep = p.get("sub_rep", "")
    sub_addr = p.get("sub_addr", "")
    con_type = p.get("construction_type", "공사")
    supply = int(p.get("supply_amount", 0))
    vat = round(supply * 0.1)
    total = supply + vat

    start_raw = p.get("start_date", "")
    end_raw = p.get("end_date", "")
    contract_raw = p.get("contract_date", datetime.today().strftime("%Y-%m-%d"))
    prog2_cond = p.get("progress2_cond", "공종 작업 완료 이후")

    start_kor = format_date_kor(start_raw) if start_raw else ""
    end_kor = format_date_kor(end_raw) if end_raw else ""
    contract_sign = format_date_kor(contract_raw)

    # 대금 (30-30-30-10)
    dep = round(supply * 0.3)
    pr1 = round(supply * 0.3)
    pr2 = round(supply * 0.3)
    bal = supply - dep - pr1 - pr2

    # 연/월/일 분리 형식 (표2에서 "2026년 02 월 05 일" 형식)
    def date_spaced(raw):
        try:
            d = datetime.strptime(raw.strip(), "%Y-%m-%d")
            return f"{d.year}년 {d.month:02d} 월 {d.day:02d} 일"
        except Exception:
            return raw

    contract_spaced = date_spaced(contract_raw)

    # 공사 유형 단축 (예: 위생배관공사 → 위생공사)
    type_short_map = {
        "위생배관공사": "위생공사",
        "음향공사": "음향공사",
        "전기공사": "전기공사",
        "냉난방공사": "냉난방공사",
        "인테리어공사": "인테리어공사",
        "도장공사": "도장공사",
    }
    con_short = type_short_map.get(con_type, con_type)

    replacements = {
        # 제목
        "위생배관공사 계약서": f"{con_type} 계약서",
        # 1조
        "위생배관공사(이하, '위생공사' 또는 '공사')": f"{con_type}(이하, '{con_short}' 또는 '공사')",
        # 5조 준공일 공사명
        "\"을\"이 위생배관공사를 완성하고": f"\"을\"이 {con_type}를 완성하고",
        # 표1 을 상호
        "호림ENG": sub_name,
        # 표2 계약금액 (VAT 별도)
        "￦48,000,000 (VAT 별도)": f"￦{supply:,} (VAT 별도)",
        # 표2 착공~준공
        "2026년 02월 05일 ~ 2026년 03월 16일": f"{start_kor} ~ {end_kor}",
        # 표3 착공/준공
        "(1) | 착공년월일 | : 2026년 02월 05일": f"(1) | 착공년월일 | : {start_kor}",  # fallback용
        ": 2026년 02월 05일": f": {start_kor}",
        ": 2026년 03월 16일": f": {end_kor}",
        # 표4 금액
        "￦ 52,800,000원": f"￦ {total:,}원",
        "￦ 48,000,000원": f"￦ {supply:,}원",
        "￦ 4,800,000원": f"￦ {vat:,}원",
        # 표5 대금 지급 (계약금30% = 중도금1차30% = 중도금2차30%, 잔금10%)
        # 각 항목이 별도 단락이므로 금액과 조건을 따로 치환
        ": ￦ 14,400,000원": f": ￦ {dep:,}원",  # 계약금/중도금1/중도금2 모두 동일(30%)
        "(공종 작업 완료 이후)": f"({prog2_cond})",  # 중도금2차 조건만 변경 가능
        ": ￦ 4,800,000원": f": ￦ {bal:,}원",  # 잔금(10%) — 부가세(10%)와 동일 금액
        # 계약 내용 기타사항 (갑/을 상호 포함 문장)
        f'㈜글로우서울(이하 \u201c갑\u201d)과 호림ENG(이하 \u201c을\u201d)는': f'㈜글로우서울(이하 \u201c갑\u201d)과 {sub_name}(이하 \u201c을\u201d)는',
        # 계약일
        "2026년 02 월 05 일": contract_spaced,
        # 서명란 갑
        "서울특별시 용산구 이태원동 164-2": GAP_INFO["addr_sign"],
        # 서명란 을
        "이 만 식  (인)": f"{sub_rep}  (인)",
        "서울특별시 강동구 구천면로 361, 102동 601호(암사동, 암사 THE Nest)": sub_addr,
    }

    doc = Document(CONSTRUCTION_TEMPLATE)
    changed = replace_in_doc(doc, replacements)
    logger.info(f"공사 계약서: {changed}개 항목 치환 → {output_path}")

    out = pathlib.Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)


# ── 일괄 생성 (Excel) ────────────────────────────────────────────────────────

def generate_from_excel(excel_path: str, output_dir: str) -> list[dict]:
    """
    Excel 파일에서 계약서 일괄 생성.

    Excel 시트:
      Sheet "자재납품계약서": 자재납품 계약서 목록
      Sheet "공사계약서":     공사 계약서 목록

    Returns:
      [{"file": 경로, "status": "ok"/"error", "msg": ...}, ...]
    """
    import openpyxl

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    out_dir = pathlib.Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    results = []

    # ── 자재납품 시트 ──
    if "자재납품계약서" in wb.sheetnames:
        ws = wb["자재납품계약서"]
        headers = [c.value for c in ws[1]]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            data = dict(zip(headers, row))
            # None → ""
            data = {k: ("" if v is None else str(v).strip()) for k, v in data.items()}
            sub = data.get("협력사명", "협력사")
            mat = data.get("자재유형", "자재")
            cdate = data.get("계약일", datetime.today().strftime("%Y-%m-%d"))
            cdate_fmt = cdate.replace("-", "")
            fname = f"{cdate_fmt}_자재납품_{mat}_{sub}.docx"
            out_path = str(out_dir / fname)
            try:
                params = {
                    "sub_name": sub,
                    "sub_name_full": data.get("협력사_법인명", sub),
                    "sub_rep": data.get("대표자", ""),
                    "sub_addr": data.get("주소", ""),
                    "sub_biz_no": data.get("사업자번호", ""),
                    "material_type": mat,
                    "material_desc": data.get("자재상세설명", f"{mat} 납품 (견적서 참조)"),
                    "total_amount": int(float(data.get("계약금액", "0").replace(",", "") or 0)),
                    "deposit_pct": int(data.get("계약금비율(%)", "40") or 40),
                    "progress_pct": int(data.get("중도금비율(%)", "50") or 50),
                    "delivery_date": data.get("납품일자", ""),
                    "install_date": data.get("설치완료일", ""),
                    "contract_date": cdate,
                }
                generate_material_contract(params, out_path)
                results.append({"file": fname, "status": "ok", "msg": "생성 완료"})
            except Exception as e:
                results.append({"file": fname, "status": "error", "msg": str(e)})

    # ── 공사 시트 ──
    if "공사계약서" in wb.sheetnames:
        ws = wb["공사계약서"]
        headers = [c.value for c in ws[1]]
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            data = dict(zip(headers, row))
            data = {k: ("" if v is None else str(v).strip()) for k, v in data.items()}
            sub = data.get("협력사명", "협력사")
            con = data.get("공사유형", "공사")
            cdate = data.get("계약일", datetime.today().strftime("%Y-%m-%d"))
            cdate_fmt = cdate.replace("-", "")
            fname = f"{cdate_fmt}_공사_{con}_{sub}.docx"
            out_path = str(out_dir / fname)
            try:
                params = {
                    "sub_name": sub,
                    "sub_rep": data.get("대표자", ""),
                    "sub_addr": data.get("주소", ""),
                    "construction_type": con,
                    "supply_amount": int(float(data.get("공급가액", "0").replace(",", "") or 0)),
                    "start_date": data.get("착공일", ""),
                    "end_date": data.get("준공예정일", ""),
                    "progress2_cond": data.get("중도금2차조건", "공종 작업 완료 이후"),
                    "contract_date": cdate,
                }
                generate_construction_contract(params, out_path)
                results.append({"file": fname, "status": "ok", "msg": "생성 완료"})
            except Exception as e:
                results.append({"file": fname, "status": "error", "msg": str(e)})

    return results
